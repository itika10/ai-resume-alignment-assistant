from io import BytesIO
from docx import Document


def add_heading(doc, text, level=1):
    if text:
        doc.add_heading(text, level=level)


def add_bullet_list(doc, items):
    for item in items:
        if item:
            doc.add_paragraph(item, style="List Bullet")


def generate_resume_docx(resume_data: dict) -> BytesIO:
    doc = Document()

    # Header
    if resume_data.get("name"):
        doc.add_heading(resume_data["name"], level=0)

    contact_parts = []
    if resume_data.get("email"):
        contact_parts.append(resume_data["email"])
    if resume_data.get("phone"):
        contact_parts.append(resume_data["phone"])

    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))

    # Summary
    if resume_data.get("summary"):
        add_heading(doc, "Professional Summary", level=1)
        doc.add_paragraph(resume_data["summary"])

    # Skills
    skill_categories = resume_data.get("skill_categories") or []
    flat_skills = resume_data.get("skills") or []

    if skill_categories or flat_skills:
        add_heading(doc, "Skills", level=1)

        if skill_categories:
            for cat in skill_categories:
                category_name = (cat.get("category") or "").strip()
                items = [str(i).strip() for i in (cat.get("items") or []) if str(i).strip()]
                if not items:
                    continue

                paragraph = doc.add_paragraph()
                if category_name:
                    label_run = paragraph.add_run(f"{category_name}: ")
                    label_run.bold = True
                paragraph.add_run(", ".join(items))
        else:
            # Fallback: render the flat list as a single paragraph
            doc.add_paragraph(", ".join(flat_skills))

    # Experience
    if resume_data.get("experience"):
        add_heading(doc, "Experience", level=1)
        for exp in resume_data["experience"]:
            role = exp.get("role", "")
            company = exp.get("company", "")
            title_line = " | ".join(part for part in [role, company] if part)
            if title_line:
                doc.add_paragraph(title_line, style="Heading 2")
            add_bullet_list(doc, exp.get("bullets", []))

    # Projects
    if resume_data.get("projects"):
        add_heading(doc, "Projects", level=1)
        for proj in resume_data["projects"]:
            title = (proj.get("title") or "").strip()
            description = (proj.get("description") or "").strip()
            start_date = (proj.get("start_date") or "").strip()
            end_date = (proj.get("end_date") or "").strip()

            if title:
                doc.add_paragraph(title, style="Heading 2")

            date_range = " – ".join(p for p in [start_date, end_date] if p)
            if date_range:
                date_para = doc.add_paragraph()
                date_run = date_para.add_run(date_range)
                date_run.italic = True

            if description:
                doc.add_paragraph(description)

            add_bullet_list(doc, proj.get("bullets", []))

            tech_stack = [str(t).strip() for t in (proj.get("tech_stack") or []) if str(t).strip()]
            if tech_stack:
                tech_para = doc.add_paragraph()
                label_run = tech_para.add_run("Tech Stack: ")
                label_run.bold = True
                tech_para.add_run(", ".join(tech_stack))

    # Certifications
    if resume_data.get("certifications"):
        add_heading(doc, "Certifications", level=1)
        add_bullet_list(doc, resume_data["certifications"])

    # Education
    if resume_data.get("education"):
        add_heading(doc, "Education", level=1)
        for edu in resume_data["education"]:
            degree = (edu.get("degree") or "").strip()
            area = (edu.get("area") or "").strip()
            institution = (edu.get("institution") or "").strip()
            location = (edu.get("location") or "").strip()
            start_date = (edu.get("start_date") or "").strip()
            end_date = (edu.get("end_date") or "").strip()

            paragraph = doc.add_paragraph()

            head_parts = []
            if degree and area:
                head_parts.append(f"{degree}, {area}")
            elif degree:
                head_parts.append(degree)
            elif area:
                head_parts.append(area)

            if head_parts:
                head_run = paragraph.add_run(head_parts[0])
                head_run.bold = True

            if institution:
                if head_parts:
                    paragraph.add_run(f" — {institution}")
                else:
                    inst_run = paragraph.add_run(institution)
                    inst_run.bold = True

            meta_parts = [p for p in [location] if p]
            date_range = " – ".join(p for p in [start_date, end_date] if p)
            if date_range:
                meta_parts.append(date_range)
            if meta_parts:
                meta_para = doc.add_paragraph()
                meta_run = meta_para.add_run(" · ".join(meta_parts))
                meta_run.italic = True

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output